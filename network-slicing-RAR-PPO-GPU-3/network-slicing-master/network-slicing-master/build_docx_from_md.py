from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "paper_draft_content.md"
OUT = ROOT / "paper_draft_v2.docx"


def set_run_font(run, size=12, bold=False, east_asia="宋体", latin="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = latin
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def add_heading_paragraph(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, east_asia="黑体")
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, size=14, bold=True, east_asia="黑体")
    else:
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, east_asia="黑体")


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    for i, line in enumerate(text.splitlines()):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, size=12, bold=False)


def main():
    content = SRC.read_text(encoding="utf-8")
    blocks = [block.strip() for block in content.split("\n\n") if block.strip()]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for block in blocks:
        if block.startswith("# "):
            add_heading_paragraph(doc, block[2:].strip(), 1)
        elif block.startswith("## "):
            add_heading_paragraph(doc, block[3:].strip(), 2)
        elif block.startswith("### "):
            add_heading_paragraph(doc, block[4:].strip(), 3)
        else:
            add_body_paragraph(doc, block)

    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        fallback = ROOT / "paper_draft_v3.docx"
        doc.save(fallback)
        print(fallback)


if __name__ == "__main__":
    main()
